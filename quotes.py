import requests
from bs4 import BeautifulSoup
from docx import Document

# URL of the website to scrape
url = 'http://quotes.toscrape.com/'

# Send a GET request to the website
response = requests.get(url)

# Parse the HTML content using BeautifulSoup
soup = BeautifulSoup(response.content, 'html.parser')

# Find all quote containers with class="quote"
quote_containers = soup.find_all('div', class_='quote')

# Initialize a Word document
doc = Document()

# Extract quotes and authors
for container in quote_containers:
    quote_text = container.find('span', class_='text').text.strip()
    author_tag = container.find('small', class_='author', itemprop='author')
    
    # Check if the author element exists in the container
    if author_tag:
        author = author_tag.text.strip()
        # Add the quote and author to the Word document
        doc.add_paragraph(f'"{quote_text}" - {author}\n\n')
    else:
        # Handle the case where the author is missing
        doc.add_paragraph(f'"{quote_text}" - Unknown Author\n\n')

# Save the Word document
doc.save('quotess.docx')

print("Quotes extracted and saved to quotes.docx")
